# pip install tkinterdnd2 Pillow PyMuPDF python-docx ebooklib reportlab bs4

import os
import sys
import subprocess
import tkinter as tk
from tkinter import ttk, messagebox
from tkinterdnd2 import DND_FILES, TkinterDnD
from PIL import Image  # Only used here for consistency with the original script
import fitz  # PyMuPDF for PDF text extraction
from docx import Document  # For reading/writing DOCX
from ebooklib import epub  # For reading/writing EPUB
from bs4 import BeautifulSoup  # For parsing EPUB html content
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

SUPPORTED_INPUTS = ['txt', 'pdf', 'doc', 'docx', 'epub']
TARGET_FORMATS = ['txt', 'pdf', 'docx', 'epub']

def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()[1:]
    text = ""
    try:
        if ext == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        elif ext == 'pdf':
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text()
        elif ext in ['doc', 'docx']:
            doc = Document(file_path)
            text = "\n".join(para.text for para in doc.paragraphs)
        elif ext == 'epub':
            book = epub.read_epub(file_path)
            for item in book.get_items():
                if item.get_type() == epub.ITEM_DOCUMENT:
                    soup = BeautifulSoup(item.get_content(), "html.parser")
                    text += soup.get_text() + "\n"
        else:
            messagebox.showerror("Unsupported Format", f"Format {ext} is not supported for extraction.")
    except Exception as e:
        messagebox.showerror("Extraction Error", f"Error extracting text from {os.path.basename(file_path)}: {str(e)}")
    return text

def write_txt(text, output_path):
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)

def write_pdf(text, output_path):
    c = canvas.Canvas(output_path, pagesize=letter)
    width, height = letter
    lines = text.splitlines()
    y = height - 40
    for line in lines:
        if y < 40:
            c.showPage()
            y = height - 40
        c.drawString(40, y, line)
        y -= 15
    c.save()

def write_docx(text, output_path):
    doc = Document()
    for para in text.splitlines():
        doc.add_paragraph(para)
    doc.save(output_path)

def write_epub(text, output_path):
    book = epub.EpubBook()
    book.set_identifier("id123456")
    book.set_title("Converted Document")
    book.set_language("en")
    
    chapter = epub.EpubHtml(title='Chapter 1', file_name='chap_01.xhtml', lang='en')
    chapter.content = f'<html><body><pre>{text}</pre></body></html>'
    book.add_item(chapter)
    
    book.toc = (epub.Link('chap_01.xhtml', 'Chapter 1', 'chap_01'),)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ['nav', chapter]
    epub.write_epub(output_path, book)

class FileConverterApp:
    def __init__(self, root):
        self.root = root
        self.root.title("File Converter")
        self.root.geometry("300x450")
        self.folder_path = os.path.join(os.path.expanduser("~"), "Downloads")
        self.loaded_files = []
        self.input_formats = SUPPORTED_INPUTS
        self.target_formats = TARGET_FORMATS

        self.setup_ui()

    def setup_ui(self):
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        main_frame.columnconfigure(0, weight=1)

        # Drop zone
        self.drop_frame = ttk.LabelFrame(main_frame, text="Drop Files Here", padding="20")
        self.drop_frame.grid(row=0, column=0, sticky="nsew", pady=10)
        self.drop_frame.columnconfigure(0, weight=1)
        self.drop_frame.rowconfigure(0, weight=1)

        self.drop_label = ttk.Label(self.drop_frame, text="Drag and drop files here")
        self.drop_label.grid(row=0, column=0, pady=50)
        self.drop_frame.drop_target_register(DND_FILES)
        self.drop_frame.dnd_bind('<<Drop>>', self.drop)

        # Remove Duplicates button, initially hidden
        self.remove_duplicates_button = ttk.Button(self.drop_frame, text="Remove Duplicates", command=self.remove_duplicates)
        self.remove_duplicates_button.grid(row=1, column=0, pady=10)
        self.remove_duplicates_button.grid_remove()  # Hide initially

        # Clear files button
        self.clear_loaded_files_button = ttk.Button(self.drop_frame, text="Clear", command=self.clear_loaded_files)
        self.clear_loaded_files_button.grid(row=2, column=0, pady=10)
        self.clear_loaded_files_button.grid_remove()  # Hide initially

        # Format selection
        ttk.Label(main_frame, text="Convert to:").grid(row=1, column=0, pady=10)
        self.target_format = tk.StringVar(value=self.target_formats[0])
        format_menu = ttk.OptionMenu(main_frame, self.target_format, self.target_formats[0], *self.target_formats)
        format_menu.grid(row=2, column=0, pady=5)

        # Progress bar
        self.progress_bar = ttk.Progressbar(main_frame, mode='determinate')
        self.progress_bar.grid(row=3, column=0, sticky="ew", pady=10)
        self.progress_bar.grid_remove()

        self.convert_button = ttk.Button(main_frame, text="Convert", command=self.convert_files)
        self.convert_button.grid(row=4, column=0, pady=10)

        # Button to open results folder
        self.go_to_explorer_button = ttk.Button(main_frame, text="See Results", command=self.open_conversion_folder)
        self.go_to_explorer_button.grid(row=5, column=0, pady=10)
        self.go_to_explorer_button.grid_remove()  # Hide initially

        self.status_label = ttk.Label(main_frame, text="...", font=("Arial", 7, "bold"))
        self.status_label.grid(row=6, column=0, pady=5)
        self.status_label.grid_remove()  # Hide initially

    def drop(self, event):
        files = self.root.tk.splitlist(event.data)
        valid_files = [f for f in files if os.path.splitext(f)[1].lower()[1:] in self.input_formats]

        if len(files) != len(valid_files):
            messagebox.showwarning("Some Invalid Files", "Some files were not in a supported format and were excluded.")
        
        self.loaded_files.extend(valid_files)
        unique_files = list(set(self.loaded_files))

        if self.loaded_files:
            self.clear_loaded_files_button.grid()
        if len(unique_files) < len(self.loaded_files):
            self.remove_duplicates_button.grid()

        self.drop_label.configure(text=f"{len(self.loaded_files)} file(s) loaded")
        self.root.update_idletasks()

    def remove_duplicates(self):
        self.loaded_files = list(set(self.loaded_files))
        self.remove_duplicates_button.grid_remove()
        self.drop_label.configure(text=f"{len(self.loaded_files)} file(s) loaded")

    def clear_loaded_files(self):
        self.loaded_files = []
        self.clear_loaded_files_button.grid_remove()
        self.drop_label.configure(text="0 file(s) loaded")
        self.remove_duplicates()
        self.go_to_explorer_button.grid_remove()
        self.progress_bar.grid_remove()
        self.status_label.grid_remove()
        self.convert_button.grid()

    def open_conversion_folder(self):
        output_dir = os.path.join(self.folder_path, "converted_files")
        try:
            if os.name == 'nt':  # Windows
                os.startfile(output_dir)
            elif os.name == 'posix':
                if sys.platform == "darwin":  # macOS
                    subprocess.Popen(["open", output_dir])
                else:  # Linux
                    subprocess.Popen(["xdg-open", output_dir])
        except Exception as e:
            messagebox.showerror("Error", f"Error opening folder: {str(e)}")

    def convert_files(self):
        if not self.loaded_files:
            messagebox.showwarning("No Files", "No files were loaded to convert.")
            return

        total_files = len(self.loaded_files)
        self.progress_bar['maximum'] = total_files
        self.progress_bar.grid()
        output_dir = os.path.join(self.folder_path, "converted_files")
        os.makedirs(output_dir, exist_ok=True)
        self.status_label.grid()

        target = self.target_format.get().lower()
        for i, file_path in enumerate(self.loaded_files, 1):
            try:
                self.status_label['text'] = f"Converting {os.path.basename(file_path)}..."
                self.root.update_idletasks()
                text = extract_text(file_path)
                base_name = os.path.splitext(os.path.basename(file_path))[0]
                output_path = os.path.join(output_dir, f"{base_name}.{target}")
                # Write out using the chosen target format
                if target == 'txt':
                    write_txt(text, output_path)
                elif target == 'pdf':
                    write_pdf(text, output_path)
                elif target == 'docx':
                    write_docx(text, output_path)
                elif target == 'epub':
                    write_epub(text, output_path)
                else:
                    messagebox.showerror("Unsupported Target", f"Target format {target} is not supported.")
                    continue

                self.progress_bar['value'] = i
                self.root.update_idletasks()
            except Exception as e:
                messagebox.showerror("Conversion Error", f"Error converting {os.path.basename(file_path)}: {str(e)}")

        self.status_label['text'] = f"Done! Files saved in {output_dir}"
        self.clear_loaded_files_button.configure(text="Restart Converter")
        self.remove_duplicates()
        self.go_to_explorer_button.grid()
        self.convert_button.grid_remove()

if __name__ == "__main__":
    root = TkinterDnD.Tk()
    app = FileConverterApp(root)
    root.mainloop()
